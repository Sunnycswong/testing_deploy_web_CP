#%%
## Import Library
import os
import json
import copy
import re
import string
import openai
from langchain.prompts import PromptTemplate
from langchain.chat_models import AzureChatOpenAI
from langchain.chains import LLMChain
# output as WORD file
import json
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor


#%%
index_name = "credit-proposal"
search_service = "gptdemosearch"
search_api_key = "PcAZcXbX2hJsxMYExc2SnkMFO0D94p7Zw3Qzeu5WjYAzSeDMuR5O"
storage_service = "creditproposal"
storage_api_key = "hJ2qb//J1I1KmVeDHBpwEpnwluoJzm+b6puc5h7k+dnDSFQ0oxuh1qBz+qPB/ZT7gZvGufwRbUrN+ASto6JOCw=="
connect_str = f"DefaultEndpointsProtocol=https;AccountName={storage_service};AccountKey={storage_api_key}"

doc_intell_endpoint = "https://doc-intelligence-test.cognitiveservices.azure.com/"
doc_intell_key = "9fac3bb92b3c4ef292c20df9641c7374"

# set up openai environment
os.environ["OPENAI_API_TYPE"] = "azure"
os.environ["OPENAI_API_BASE"] = "https://pwcjay.openai.azure.com/"
os.environ["OPENAI_API_VERSION"] = "2023-05-15"
os.environ["OPENAI_API_KEY"] = "f282a661571f45a0bdfdcd295ac808e7"

os.environ["AZURE_COGNITIVE_SEARCH_SERVICE_NAME"] = search_service
os.environ["AZURE_COGNITIVE_SEARCH_API_KEY"] = search_api_key
os.environ["AZURE_INDEX_NAME"] = index_name

#%%
def load_json(json_path):
    with open(json_path, "r" ,encoding="utf-8") as f:
        data = json.load(f)
    return data

#%%
rm_note_file_name = "GOGOVAN_RM_note.json"
hierarchy_file_name = "hierarchy.json"
schema_file_name = "schema.json"

rm_note_txt = load_json(rm_note_file_name)[0]
hierarchy_dict_list = load_json(hierarchy_file_name)
schema_dict_list = load_json(schema_file_name)
hierarchy_rm_note_list = load_json("GOGOVAN_hierarchy_rm_note.json")

#%%

def cap(match):
    return(match.group().capitalize())

#%%
def update_schema_json(schema_dict_list, hierarchy_dict_list, update_key='RM Note'):
    """
    This function will update the schema_dict_list (deepcopy) JSON by the specified key (e.g. RM Note).
    You will update the values based on hierarchy_dict_list (eitehr stored the RM info or Document info)
    """
    updated_schema_dict_list = copy.deepcopy(schema_dict_list)
    update_dict = {}
    for l in hierarchy_dict_list:
        section_name = l["Section"]
        if section_name not in update_dict:
            update_dict[section_name] = ""
        if len(l['Value']) > 0:
            if l['Breakdown'] == "":
                update_dict[section_name] += "- {}: {}\n".format(l['Sub Section'], l['Value'])
            else:
                update_dict[section_name] += "- {} ({}): {}\n".format(l['Sub Section'], l['Breakdown'], l['Value'])
    print(update_dict)
    for schema_dict in updated_schema_dict_list:
        if schema_dict['Section'] in update_dict.keys():
            if update_key in schema_dict:
                schema_dict[update_key] = update_dict[schema_dict['Section']]
            else:
                raise ValueError("You do not have key value: {} in your schema_dict_list JSON".format(update_key))
    return updated_schema_dict_list

schema_with_rm_note = update_schema_json(schema_dict_list, hierarchy_rm_note_list, update_key='RM Note')
json.dump(schema_with_rm_note, open("GOGOVAN_schema_with_extraction.json", "w"), indent=4)


#The way to embed and combine the document extraction with RM extraction
# combine the document extraction and RM extraction result for proposal writing

# hierarchy_document
hierarchy_document_list = load_json("hierarchy.json")
schema_with_both_extract = update_schema_json(schema_with_rm_note, hierarchy_document_list, update_key='Document')
json.dump(schema_with_both_extract, open("GOGOVAN_schema_with_extraction.json", "w"), indent=4)


#%%
def proposal_proposal_template_text():
    proposal_proposal_template_text = """
        Read the input prompt, RM note (Keyword:rm_note), rm_note (Keyword: rm_note), document (Keyword: document), and example for this section carefully:
        
        Prompt:
        {prompt}
        ======
        RM Note: (Keyword: rm_note)
        {rm_note}
        ======
        Document: (Keyword: document)
        {document}
        ======
        Example: (Keyword: proposal_example)
        {example}
        ======
        
        Then write paragraph(s) based on the above aggregrated context

        Rules you need to follow:
        1. Don't mention the word "RM Note" and "Component", and don't mention you held a meeting with the client! Instead, you shall say "It is mentioned that"
        2. Don't mention the source of your input (i.e. RM Note (Keyword: rm_note), example, document)
        3. Don't justify your answers
        4. Don't provide suggestion or recommendation by yourself
        5. Provide your answer in English
        6. Breake it to multi-paragraphs if one single paragraph consists of more than 100 words
        7. In the same paragraph, don't input line breaks among the sentences
        8. Don't start with your answer by a title. You must start your paragraph immediately
        9. The example (Keyword: proposal_example) above is just for your reference only to improve your theme, you must not directly copy the content in the examples
        10. If possible, you can use point-form, tables to provide your answer
        11. Don't introduce what this section includes

        Guidance when you do not have the information:
        1. When you don't have the specific information or you need further information (Keyword: further_info), you have to write it in the following format: [RM please helps provide the further information of (Keyword: further_info)], where please supplement the information you need here.
        2. You must not create the information by yourself if you don't have relevant information
        3. You cannot say "It's unclear that", please refer to point 1 for the formatting for requesting further information

        Take a deep breath and work on this step by step
        """
    prompt_template_proposal = PromptTemplate(template=proposal_proposal_template_text, input_variables=["prompt", "rm_note", "document", "example"])

    llm_proposal = AzureChatOpenAI(deployment_name="gpt-35-16k", temperature=0,
                            openai_api_version="2023-05-15", openai_api_base="https://pwcjay.openai.azure.com/")
    
    return prompt_template_proposal, llm_proposal

def core_llm_call(prompt, rm_note, document, example):
    prompt_template_proposal, llm_proposal = proposal_proposal_template_text()

    """
    A core function to generate the proposal per section

    Parameters:
    -----------
    prompt: str
    Prompt text for instructing the output based on RM prompt

    rm_note: str
    It contains the information from the RM

    example: str
    Input example for GPt to take it as an example

    """
    #.format(content=content, section=section, context=context)
    
    chain = LLMChain(
        llm=llm_proposal,
        prompt=prompt_template_proposal
    )
    drafted_text = chain({"prompt": prompt
                    , "rm_note":rm_note
                    , "document":document
                    , "example":example})['text']
    drafted_text2 = drafted_text.replace("Based on the given information, ", "").replace("It is mentioned that ", "")
    
    #All capital letters for first letter in sentences
    formatter = re.compile(r'(?<=[\.\?!]\s)(\w+)')
    drafted_text2 = formatter.sub(cap, drafted_text2)

    #output the result
    return drafted_text2


#%%

def gen_first_web_and_json(schema_with_both_extract):
    old_schema = copy.deepcopy(schema_with_both_extract) # your old schema data
    new_schema = [
        {
            "Section": "Project Overview",
            "Component": [
                "Project Name",
                "Project Description",
                "Project Cost"
            ],
            "Prompt": "Write me an overview on the project based on the provided component",
            "Example": "",
            "RM Note": "",
            "Document": "",
            "GPT Output": "",
            "Component_Text": "Project Name, Project Description, Project Cost"
        },
        {
            "Section": "Client Background",
            "Component": [],
            "Prompt": "Write the client background section by the given information",
            "Example": "",
            "RM Note": "",
            "Document": "",
            "GPT Output": "",
            "Component_Text": ""
        },
        {
            "Section": "Credit Request" ,
            "Component": [
                "Relationship with BEA",
                "Amount Requested",
                "Credit Facility",
                "Purpose of Financing",
                "Loan Term",
                "Repayment Plan",
                "Project Details"
            ],
            "Prompt": "Write the client background section by the given information",
            "Example": "",
            "RM Note": "",
            "Document": "",
            "GPT Output": "",
            "Component_Text": "Relationship with BEA, Amount Requested, Credit Facility, Purpose of Financing, Loan Term, Repayment Plan"
        },
        {
            "Section": "Collateral and Guarantees",
            "Component": [
                "Financial Statements",
                "Projections",
                "Personal Guarantees",
                "Corporate Guarantees"
            ],
            "Prompt": "Write the client background section by the given information",
            "Example": "",
            "RM Note": "",
            "Document": "",
            "GPT Output": "",
            "Component_Text": "Financial Statements, Projections, Personal Guarantees, Corporate Guarantees"
        },
        {
            "Section": "Conclusion",
            "Component": [
                "Collateral and Guarantees",
                "Risks and Mitigation",
                "Compliance and Legal",
                "Ongoing Monitoring",
                "Regular Legal Reviews and Compliance Assessment",
                "Risk Mitigation",
                "Collateral Valuation",
                "Overall Recommendation"
            ],
            "Prompt": "Write the client background section by the given information",
            "Example": "",
            "RM Note": "",
            "Document": "",
            "GPT Output": "",
            "Component_Text": "Collateral and Guarantees, Risks and Mitigation, Compliance and Legal"
        }
    ]  # initialize the new schema data

    # Create a mapping for the old sections to the new sections
    section_mapping = {
        "A. Project Overview": "Project Overview",
        "B. Client Background": "Client Background",
        "C. Credit Request": "Credit Request",
        "D. Project Details": "Credit Request",
        "E. Financial Information": "Collateral and Guarantees",
        "F. Collateral and Guarantees": "Collateral and Guarantees",
        "G. Repayment Capacity": "Conclusion",
        "H. Risk and Mitigation": "Conclusion",
        "I. Compliance and Legal": "Conclusion",
        "J. Conclusion": "Conclusion",
        "K. Recommendations": "Conclusion",
        "Final Approval": "Conclusion"
    }

    # Initialize an empty dict to store the new schema
    new_schema_dict = {}

    # Iterate over the old schema
    for item in old_schema:
        # Get the old section name
        old_section = item["Section"]
        # Get the new section name from the mapping
        new_section = section_mapping[old_section]

        # If the new section is already in the new schema, append the components
        if new_section in new_schema_dict:
            new_schema_dict[new_section]["Component"].extend(item["Component"])
            new_schema_dict[new_section]["Component_Text"] += ", " + item["Component_Text"]
        else:
            # Otherwise, create a new section with the same data as the old one
            new_schema_dict[new_section] = item
            new_schema_dict[new_section]["Section"] = new_section

    # Convert the new schema dict back into a list
    new_schema = list(new_schema_dict.values())

    json.dump(new_schema, open("testing_schema.json", "w"), indent=4)


    # run with GPT by sectors
    New_schema_dict_list_with_GPT = copy.deepcopy(new_schema)
    for l in New_schema_dict_list_with_GPT:
        section = l["Section"]
        prompt = l['Prompt']
        rm_note = l['RM Note']
        document = l['Document']
        example = l['Example']
        l["GPT Output"] = core_llm_call(prompt, rm_note, document, example)
        print(section)
        print(l["GPT Output"])
        print("="*20)

    json.dump(New_schema_dict_list_with_GPT, open("testing_schema.json", "w"), indent=4)

    return New_schema_dict_list_with_GPT

gen_first_web_and_json(schema_with_both_extract)


#%%
def edit_web_json(section,prompt):
    document_list = load_json("testing_schema.json")
    for i in document_list:
        # Edit 'Project Overview'
        if section == i['Section']:
            rm_note = i['RM Note']
            document = i['Document']
            example = i['Example']
            i["GPT Output"] = core_llm_call(prompt, rm_note, document, example)
            print(i["GPT Output"])
            print_sector = i 
    
    json.dump(document_list, open("testing_schema.json", "w"), indent=4)

    return print_sector

# edit_web_json('Client Background',"Gogovan is requesting a credit facility of $60 million instead of $10 million")


#%%

# Old schema.
"""schema_dict_list_with_GPT = copy.deepcopy(schema_with_both_extract)
for l in schema_dict_list_with_GPT:
    section = l["Section"]
    prompt = l['Prompt']
    rm_note = l['RM Note']
    document = l['Document']
    example = l['Example']
    l["GPT Output"] = core_llm_call(prompt, rm_note, document, example, llm_proposal, prompt_template_proposal)
    print(section)
    print(l["GPT Output"])
    print("="*20)

#%%
json.dump(schema_dict_list_with_GPT, open("GOGOVAN_schema_with_extraction_GPT.json", "w"), indent=4)
"""


#%%

def create_docx():

    # Load the JSON data
    with open('testing_schema.json', 'r') as json_file:
        json_data = json.load(json_file)

    # Create a new Word document
    document = Document()

    # extract client name for hierarchy_document
    hierarchy_document = load_json("hierarchy.json")
    rm_note_txt = load_json(rm_note_file_name)[0]
    client_name = ""
    if str(hierarchy_document[0]["Value"]) == "":
        lines = rm_note_txt.split('\n')
        # First check whether there is a ': ' in the string
        for line in lines:
            if line.startswith('Client:'):
                client_name = line.split('Client: ')[1]
    else:
        client_name = str(hierarchy_document[0]["Value"])

    title_text = "Credit Proposal for ", client_name
    title_size = 20 # Font size in points

    # Create a paragraph for the title
    title_paragraph = document.add_paragraph()

    # Add the title text to the paragraph
    title_run = title_paragraph.add_run(title_text)

    # Apply formatting to the title run
    title_run.bold = True
    title_run.font.size = Pt(title_size)

    # Set the alignment of the paragraph to the center
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Convert JSON values to section headers and paragraphs in the Word document
    for item in json_data:
        section = item['Section']
        context = item['GPT Output']

        # for the part of Financial Information, attach the STATEMENTs as table 
        if section == "5. Financial Information":
            # Add the section header
            document.add_heading(section, level=1)
            # Add the context as a paragraph
            document.add_paragraph(context)

            # Add table
            table_title = document.add_paragraph("")
            table_title.add_run('CONSOLIDATED STATEMENT OF PROFIT OR LOSS for the year ended 31 December in HK$ million').bold = True
            # Define the table data
            table_data = [
                ['for the year ended 31 December in HK$ million', 'Note', '2022', '2021'],
                ['Revenue from Hong Kong transport operations', '4', '13,404', '13,177'],
                ['Revenue from Hong Kong station commercial businesses', '5', '3,077', '3,208'],
                ['Revenue from Hong Kong property rental and management businesses', '6', '4,779', '5,036'],
                ['Revenue from Mainland China and international railway, property rental and management subsidiaries', '7', '26,016', '25,045'],
                ['Revenue from other businesses', '8', '363', '383'],
                ['Revenue from Mainland China property development', '7', '173', '353'],
                ['Total revenue', '', '47,812', '47,202'],
                ['Expenses relating to Hong Kong transport operations - Staff costs and related expenses', '10A', '(6,341)', '(6,155)'],
                ['Expenses relating to Hong Kong transport operations - Maintenance and related works', '', '(2,221)', '(2,339)'],
                ['Expenses relating to Hong Kong transport operations - Energy and utilities', '', '(1,991)', '(1,801)'],
                ['Earnings per share - Basic', '18', 'HK$1.59', 'HK$1.55'],
                ['Earnings per share - Diluted', '18', 'HK$1.59', 'HK$1.54']
            ]

            # Create a table with the table data
            table = document.add_table(rows=1, cols=4)
            table.style = 'Table Grid'  # Apply table style


            table.allow_autofit = False
            table.cell(0,0).width = Inches(1.0)                                                
            table.cell(0,1).width = Inches(1.0)

            # Add the table header row
            header_row = table.rows[0]
            header_row.cells[0].text = 'for the year ended 31 December in HK$ million'
            header_row.cells[1].text = 'Note'
            header_row.cells[2].text = '2022'
            header_row.cells[3].text = '2021'

            # Add the remaining rows of the table data
            for row_data in table_data[1:]:
                row = table.add_row().cells
                row[0].text = row_data[0]
                row[1].text = row_data[1]
                row[2].text = row_data[2]
                row[3].text = row_data[3]
        else:
            # Add the section header
            document.add_heading(section, level=1)

            # Split context into lines and check each line
            for line in context.split('\n'):
                # Create a new paragraph for each line
                paragraph = document.add_paragraph()
                
                # Search for the pattern [RM please ... ] using regex
                matches = re.findall(r'\[RM please .*?\]', line)

                if matches:
                    # If there's a match, split line into parts
                    parts = re.split(r'(\[RM please .*?\])', line)

                    for part in parts:
                        run = paragraph.add_run(part)
                        
                        if part in matches:
                            # This part should be colored red
                            run.font.color.rgb = RGBColor(255, 0, 0)  # RGB values for red
                else:
                    # Normal text
                    run = paragraph.add_run(line)

    # Save the Word document
    document.save('GOGOVAN_Word_proposal.docx')

    return

#%%
create_docx()

# %%
rm_note_txt
# %%
